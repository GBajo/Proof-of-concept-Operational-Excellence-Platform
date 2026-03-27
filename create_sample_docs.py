"""
create_sample_docs.py — Genera documentos de ejemplo para el módulo RAG.

Ejecutar una vez:
    python create_sample_docs.py

Requiere:
    pip install python-docx openpyxl
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment


DOCS_DIR = Path(__file__).parent / "docs"
DOCS_DIR.mkdir(exist_ok=True)


# ─────────────────────────────────────────────────────────────
# 1. SOP de limpieza de línea de empaquetado (DOCX)
# ─────────────────────────────────────────────────────────────
def create_sop_limpieza():
    doc = Document()

    # Título
    title = doc.add_heading("", 0)
    run = title.add_run("SOP-LIN-001 — Procedimiento Estándar de Limpieza de Línea de Empaquetado")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0, 57, 112)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("Versión: 3.2  |  Fecha: 15/01/2026  |  Área: Producción / Empaquetado", style="Normal")
    doc.add_paragraph("Aprobado por: Dpto. Calidad  |  Próxima revisión: 15/01/2027", style="Normal")
    doc.add_paragraph()

    # 1. Objeto
    doc.add_heading("1. Objeto", level=1)
    doc.add_paragraph(
        "Este procedimiento establece las instrucciones detalladas para llevar a cabo la limpieza "
        "y sanitización de la línea de empaquetado de productos farmacéuticos sólidos (comprimidos "
        "y cápsulas), garantizando la ausencia de contaminación cruzada entre lotes y el cumplimiento "
        "de los requisitos de las Buenas Prácticas de Fabricación (BPF/GMP)."
    )

    # 2. Alcance
    doc.add_heading("2. Alcance", level=1)
    doc.add_paragraph(
        "Este SOP aplica a todas las líneas de empaquetado (L1–L5) de la planta de fabricación. "
        "Cubre la limpieza entre lotes del mismo producto (limpieza menor), entre productos diferentes "
        "(limpieza mayor) y la limpieza de fin de turno."
    )

    # 3. Responsabilidades
    doc.add_heading("3. Responsabilidades", level=1)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Rol"
    hdr[1].text = "Responsabilidad"
    rows = [
        ("Operario de línea", "Ejecutar la limpieza según este procedimiento y registrar en el parte de turno."),
        ("Supervisor de producción", "Verificar la ejecución y firmar el registro de limpieza."),
        ("Control de Calidad", "Tomar muestras de hisopado cuando corresponda y liberar la línea."),
        ("Mantenimiento", "Asistir en el desmontaje de piezas de difícil acceso cuando sea necesario."),
    ]
    for rol, resp in rows:
        row = table.add_row().cells
        row[0].text = rol
        row[1].text = resp
    doc.add_paragraph()

    # 4. Materiales y equipos
    doc.add_heading("4. Materiales y Equipos Necesarios", level=1)
    items = [
        "Agua purificada (PW) según especificación WFI-002",
        "Detergente alcalino Tego 51 al 2% en agua purificada",
        "Alcohol etílico al 70% para superficies de contacto con producto",
        "Paños de microfibra limpios (color azul: equipo; color rojo: suelo)",
        "Aspirador industrial con filtro HEPA certificado",
        "Escobillas de nylon de uso exclusivo por línea (identificadas con número de línea)",
        "EPIs: guantes de nitrilo, mascarilla FFP2, gafas de protección, bata",
        "Kit de hisopado para muestras de superficie (si aplica cambio de producto)",
    ]
    for item in items:
        doc.add_paragraph(item, style="List Bullet")
    doc.add_paragraph()

    # 5. Procedimiento — Limpieza menor
    doc.add_heading("5. Procedimiento", level=1)
    doc.add_heading("5.1 Limpieza Menor (entre lotes del mismo producto)", level=2)
    steps_menor = [
        ("Paso 1 — Parada de línea", "Detener la línea de empaquetado según el procedimiento SOP-LIN-000. Confirmar que no hay producto en curso. Registrar hora de parada en el parte de turno."),
        ("Paso 2 — Señalización", "Colocar cartel 'LÍNEA EN LIMPIEZA' en el acceso a la línea. Nadie puede acceder a la zona sin EPIs completos durante el proceso."),
        ("Paso 3 — Retirada de producto", "Retirar todos los blísteres, cajas y materiales de empaquetado pendientes. Depositar en contenedor identificado con el número de lote. Nunca mezclar material de lotes distintos."),
        ("Paso 4 — Aspiración", "Aspirar todas las superficies externas de los equipos: tolvas, guías, cintas transportadoras. Prestar especial atención a las zonas de acumulación de polvo bajo las máquinas."),
        ("Paso 5 — Limpieza húmeda", "Preparar solución detergente Tego 51 al 2%. Limpiar superficies con paño humedecido en dirección de arriba hacia abajo. No verter líquido directamente sobre componentes eléctricos."),
        ("Paso 6 — Aclarado", "Aclarar con paño humedecido en agua purificada. Repetir hasta ausencia de restos de detergente (verificar con tira pH: debe ser 6.5–7.5)."),
        ("Paso 7 — Sanitización", "Aplicar alcohol 70% con paño limpio en todas las superficies de contacto con producto. Dejar actuar 3 minutos antes de proceder."),
        ("Paso 8 — Secado", "Secar con paños secos de microfibra. La línea debe estar completamente seca antes de reiniciar la producción."),
        ("Paso 9 — Verificación visual", "El supervisor revisará la línea. Si se detectan restos de producto o suciedad, repetir desde el Paso 4."),
        ("Paso 10 — Registro", "Cumplimentar el Registro de Limpieza RL-LIN-001 con nombre, fecha, hora inicio, hora fin y firma. El supervisor contrafirma."),
    ]
    for titulo, texto in steps_menor:
        p = doc.add_paragraph()
        run_t = p.add_run(titulo + ": ")
        run_t.bold = True
        p.add_run(texto)

    # 5.2 Limpieza mayor
    doc.add_heading("5.2 Limpieza Mayor (cambio de producto)", level=2)
    doc.add_paragraph(
        "Además de todos los pasos de la limpieza menor, se deben realizar las siguientes acciones adicionales:"
    )
    steps_mayor = [
        "Desmontar todas las piezas desmontables según la guía GM-LIN-001 (tolvas, guías de blíster, ruedas de formato).",
        "Limpiar piezas desmontadas en la sala de lavado con solución detergente, aclarar y secar en estufa a 60°C durante 30 minutos.",
        "Tomar hisopados de superficie en los 3 puntos críticos definidos en el plan de muestreo MS-LIN-001.",
        "Esperar resultado de hisopados antes de montar las piezas. Límite de aceptación: ≤10 ppm de principio activo anterior.",
        "Registrar los cambios de formato en el Registro de Cambio de Formato CF-LIN-002.",
        "Control de Calidad debe liberar la línea mediante firma en el RL-LIN-001 antes de iniciar producción.",
    ]
    for s in steps_mayor:
        doc.add_paragraph(s, style="List Number")

    # 6. Frecuencias
    doc.add_heading("6. Frecuencias de Limpieza", level=1)
    freq_table = doc.add_table(rows=1, cols=3)
    freq_table.style = "Table Grid"
    hdr2 = freq_table.rows[0].cells
    hdr2[0].text = "Tipo de limpieza"
    hdr2[1].text = "Frecuencia"
    hdr2[2].text = "Documento de registro"
    freqs = [
        ("Limpieza de fin de turno", "Al finalizar cada turno de producción", "RL-LIN-001"),
        ("Limpieza menor entre lotes", "Cada cambio de lote del mismo producto", "RL-LIN-001"),
        ("Limpieza mayor cambio producto", "Cada cambio de producto en la línea", "RL-LIN-001 + CF-LIN-002"),
        ("Limpieza semanal profunda", "Todos los lunes antes del inicio de producción", "RL-LIN-003"),
        ("Verificación microbiológica", "Mensual en puntos críticos", "RM-LIN-004"),
    ]
    for tipo, freq, doc_ref in freqs:
        row = freq_table.add_row().cells
        row[0].text = tipo
        row[1].text = freq
        row[2].text = doc_ref

    # 7. Criterios de aceptación
    doc.add_heading("7. Criterios de Aceptación", level=1)
    doc.add_paragraph(
        "Una línea se considera limpia y lista para producción cuando se cumplen TODOS los siguientes criterios:"
    )
    criterios = [
        "Inspección visual: ausencia de polvo, restos de producto anterior, lubricantes o materiales de empaquetado.",
        "pH del agua de aclarado: entre 6.5 y 7.5 (verificado con tiras reactivas).",
        "Hisopados de superficie (solo cambio de producto): ≤10 ppm del principio activo anterior según HPLC.",
        "Ausencia de olores extraños.",
        "Documentación completa y firmada por el operario y el supervisor.",
        "Liberación formal por Control de Calidad (en cambio de producto).",
    ]
    for c in criterios:
        doc.add_paragraph(c, style="List Bullet")

    # 8. Acciones ante desviaciones
    doc.add_heading("8. Acciones ante Desviaciones", level=1)
    desviaciones = [
        ("Hisopado fuera de especificación", "Repetir limpieza mayor completa. Notificar a Control de Calidad. Abrir desviación según SOP-CAL-010."),
        ("pH de aclarado fuera de rango", "Repetir aclarado con agua purificada fresca hasta pH dentro de rango."),
        ("Daño en alguna pieza durante el desmontaje", "Notificar a Mantenimiento. No instalar la pieza dañada. Sustituir antes de producción."),
        ("Incumplimiento de tiempos de actuación del desinfectante", "Repetir aplicación de desinfectante desde el inicio del tiempo de contacto."),
    ]
    for desv, accion in desviaciones:
        p = doc.add_paragraph()
        p.add_run(f"{desv}: ").bold = True
        p.add_run(accion)

    # 9. Historial de cambios
    doc.add_heading("9. Historial de Revisiones", level=1)
    hist = doc.add_table(rows=1, cols=4)
    hist.style = "Table Grid"
    hdr3 = hist.rows[0].cells
    for i, h in enumerate(["Versión", "Fecha", "Descripción del cambio", "Aprobado por"]):
        hdr3[i].text = h
    cambios = [
        ("1.0", "10/03/2022", "Versión inicial", "J. López"),
        ("2.0", "05/09/2023", "Incorporación de limpieza mayor con hisopados. Actualización de límites.", "M. García"),
        ("3.0", "20/06/2025", "Actualización de desinfectante (cambio de proveedor). Nuevos puntos de muestreo.", "A. Martínez"),
        ("3.2", "15/01/2026", "Corrección de tiempos de contacto del alcohol 70%. Añadido paso de secado en estufa.", "A. Martínez"),
    ]
    for v, fecha, desc, aprobado in cambios:
        row = hist.add_row().cells
        row[0].text = v
        row[1].text = fecha
        row[2].text = desc
        row[3].text = aprobado

    path = DOCS_DIR / "SOP-LIN-001_Limpieza_Linea.docx"
    doc.save(str(path))
    print(f"  ✓ Creado: {path.name}")


# ─────────────────────────────────────────────────────────────
# 2. Manual de errores frecuentes de envasadora (DOCX)
# ─────────────────────────────────────────────────────────────
def create_manual_errores():
    doc = Document()

    title = doc.add_heading("", 0)
    run = title.add_run("Manual de Resolución de Errores — Envasadora Uhlmann UPS 4")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0, 57, 112)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("Modelo: Uhlmann UPS 4 (adaptable a IMA, Romaco y Marchesini de gama similar)", style="Normal")
    doc.add_paragraph("Revisión: 2.1  |  Fecha: 10/11/2025  |  Mantenimiento & Producción", style="Normal")
    doc.add_paragraph()

    doc.add_heading("Introducción", level=1)
    doc.add_paragraph(
        "Este manual recoge los 15 errores más frecuentes detectados en operación de la envasadora "
        "de blísteres, sus causas raíz identificadas, los pasos de diagnóstico y las acciones correctivas "
        "que puede ejecutar el operario de turno antes de escalar al equipo de mantenimiento. "
        "El objetivo es minimizar el tiempo de parada no planificado y el desperdicio de material."
    )

    errores = [
        {
            "codigo": "ERR-001",
            "titulo": "Rotura de lámina de aluminio en la zona de sellado",
            "sintomas": "Alarma sonora. Detector óptico de integridad activado. Blísteres salen abiertos o con sellado incompleto.",
            "causas": [
                "Temperatura de sellado superior a la especificada (nominal: 180°C ±5°C).",
                "Velocidad de línea demasiado alta para el espesor de aluminio actual.",
                "Lámina con defecto de fábrica (arrugas, golpes en el rollo).",
                "Rodillo de tracción con desgaste excesivo.",
            ],
            "acciones": [
                "Detener la línea. Retirar el material defectuoso del circuito.",
                "Verificar temperatura de los platos de sellado con termómetro de contacto. Si hay desviación >3°C, llamar a Mantenimiento.",
                "Inspeccionar el rollo de aluminio: buscar arrugas o daños. Si se detectan, cortar el tramo dañado y empalmar.",
                "Verificar tensión del rollo (no debe estar demasiado tenso ni flojo). Ajustar freno de bobina.",
                "Reducir velocidad en un 10% y observar si el problema se resuelve.",
                "Si persiste tras 5 minutos de operación a velocidad reducida, escalar a Mantenimiento con código ERR-001.",
            ],
            "tiempo_resolucion": "Operario: 10–20 min. Con Mantenimiento: 30–60 min.",
            "impacto": "Alto — produce rechazo masivo si no se detecta a tiempo.",
        },
        {
            "codigo": "ERR-002",
            "titulo": "Blíster con cavidades vacías (falta de comprimido)",
            "sintomas": "Alarma del sistema de control de presencia (CCD o vision system). Contador de rechazos aumenta.",
            "causas": [
                "Tolva de alimentación vacía o con producto apelmazado.",
                "Vibrador de alimentación desajustado o averiado.",
                "Comprimidos fuera de especificación de peso/tamaño (se atascan en guías).",
                "Guías de alimentación sucias o con acumulación de polvo.",
            ],
            "acciones": [
                "Verificar nivel de tolva. Si está vacía, rellenar con el lote correcto según la orden de fabricación.",
                "Comprobar que el vibrador esté en marcha (LED verde en panel). Si está apagado, activar desde la pantalla HMI.",
                "Limpiar guías de alimentación con brocha y aspirador. Verificar que no haya comprimidos rotos obstruyendo el paso.",
                "Pasar comprimido de muestra por las guías manualmente para verificar el flujo.",
                "Si el problema es el vibrador (no responde), notificar a Mantenimiento.",
            ],
            "tiempo_resolucion": "Operario: 5–15 min.",
            "impacto": "Alto — unidad no apta para consumo. El sistema de visión debe rechazarlos automáticamente.",
        },
        {
            "codigo": "ERR-003",
            "titulo": "Atasco en zona de troquelado / corte de blísteres",
            "sintomas": "Parada brusca de línea. Alarma mecánica. Material acumulado visible en zona de troquelado.",
            "causas": [
                "Mal alineamiento del material (PVC/aluminio) en la guía central.",
                "Cuchilla de corte desgastada o con mella.",
                "Cuerpo extraño (fragmento de comprimido, trozo de film) atrapado.",
                "Velocidad de avance descuadrada respecto al ciclo de corte.",
            ],
            "acciones": [
                "PARADA DE EMERGENCIA si hay riesgo de daño en el equipo o el operario.",
                "Retirar el material acumulado con cuidado usando las herramientas proporcionadas (nunca con las manos desnudas en la zona de cuchillas).",
                "Verificar alineación del material en la guía. Reajustar si es necesario con la ruleta de centrado.",
                "Inspeccionar cuchilla visualmente. Si hay mella visible, sustituir (solo personal de Mantenimiento).",
                "Realizar 3 ciclos en vacío (sin material) antes de reiniciar producción.",
            ],
            "tiempo_resolucion": "Operario: 5–10 min. Sustitución de cuchilla (Mantenimiento): 45 min.",
            "impacto": "Medio — parada de línea. Sin desperdicio de producto si se resuelve rápido.",
        },
        {
            "codigo": "ERR-004",
            "titulo": "Temperatura de formado fuera de rango — PVC no forma correctamente",
            "sintomas": "Blísteres con cavidades mal formadas, rotas o con burbujas. Alarma de temperatura.",
            "causas": [
                "Resistencia de calentamiento de la zona de formado averiada.",
                "Sensor de temperatura descalibrado.",
                "PVC incorrecto (diferente calibre al especificado).",
                "Flujo de aire de refrigeración de molde obstruido.",
            ],
            "acciones": [
                "Verificar temperatura en HMI: zona formado nominal 120°C ±5°C (varía según el formato).",
                "Si la temperatura real difiere >10°C del setpoint durante más de 2 minutos, detener línea y avisar a Mantenimiento.",
                "Verificar el calibre del PVC cargado. Consultar la ficha de formato para el espesor especificado.",
                "Comprobar que las rejillas de ventilación del molde no estén obstruidas.",
            ],
            "tiempo_resolucion": "Diagnóstico: 5 min. Resolución por Mantenimiento: 1–3 horas.",
            "impacto": "Alto — produce blísteres defectuosos a alta velocidad.",
        },
        {
            "codigo": "ERR-005",
            "titulo": "Impresión de código Datamatrix / número de lote defectuosa",
            "sintomas": "Sistema de visión rechaza blísteres por código ilegible o incorrecto. Alarma de impresora.",
            "causas": [
                "Cartucho de tinta de la impresora inkjet a punto de agotarse.",
                "Cabezal de impresión sucio o con boquillas obstruidas.",
                "Velocidad de línea superior a la capacidad de impresión.",
                "Parámetros de impresión (fecha de caducidad, número de lote) incorrectamente configurados.",
            ],
            "acciones": [
                "Verificar nivel de tinta en la pantalla de la impresora (Linx o Markem-Imaje). Si nivel <15%, cambiar cartucho.",
                "Realizar ciclo de limpieza de cabezal desde el menú de la impresora (Maintenance > Head Cleaning).",
                "Verificar los parámetros impresos en pantalla HMI: número de lote, fecha de fabricación y caducidad según la orden.",
                "Imprimir 5 blísteres de prueba y verificar con lector de códigos antes de reiniciar producción.",
                "Si el código sigue siendo ilegible, escalar a Mantenimiento. No producir sin serialización válida.",
            ],
            "tiempo_resolucion": "Operario: 10–15 min. Mantenimiento (fallo de cabezal): 1 hora.",
            "impacto": "Crítico — sin código válido el producto no es trazable y debe ser rechazado.",
        },
        {
            "codigo": "ERR-006",
            "titulo": "Exceso de rechazo en control de peso (IPC en línea)",
            "sintomas": "% de rechazo supera el límite de alerta (>1%) o de acción (>3%). Alarma de calidad.",
            "causas": [
                "Variabilidad de peso de los comprimidos por encima de la especificación.",
                "Báscula de control en línea descalibrada.",
                "Mezcla de dos lotes distintos en la tolva.",
                "Comprimidos rotos en la tolva por exceso de vibración.",
            ],
            "acciones": [
                "Detener el llenado de la tolva inmediatamente.",
                "Tomar muestra de 10 comprimidos de la línea y pesar manualmente. Si la variabilidad es >±5% del peso nominal, notificar a Control de Calidad.",
                "Verificar calibración de la báscula con pesas patrón certificadas (guardadas en el cajón de la báscula).",
                "Revisar tolva: si hay mezcla de lotes, vaciar y limpiar según SOP-LIN-001.",
                "Cubrir los blísteres afectados y segregar en cuarentena hasta resolución.",
            ],
            "tiempo_resolucion": "Diagnóstico: 15 min. Resolución completa: variable (puede requerir investigación de QC).",
            "impacto": "Crítico — afecta directamente a la calidad del producto.",
        },
        {
            "codigo": "ERR-007",
            "titulo": "Atasco o desbordamiento en la zona de cartoning (encajado)",
            "sintomas": "Parada de la encajadora. Blísteres acumulados en la cinta de transferencia. Alarma mecánica.",
            "causas": [
                "Velocidad desincronizada entre la blistera y la encajadora.",
                "Fallo en el sistema de inserción de folleto (folleto doblado o ausente).",
                "Cajas de cartón mal formadas o fuera de especificación.",
                "Sensor de presencia de caja en fallo.",
            ],
            "acciones": [
                "Parar ambas máquinas (blistera y encajadora) coordinadamente.",
                "Retirar los blísteres acumulados y depositarlos en bandeja de recuperación.",
                "Verificar el cargador de folletos: recargar si está vacío, retirar folletos arrugados.",
                "Verificar el cargador de cajas: separar las cajas mal troqueladas.",
                "Reiniciar la encajadora y ajustar la velocidad de sincronización desde el HMI (menú Sync).",
            ],
            "tiempo_resolucion": "Operario: 10–20 min.",
            "impacto": "Medio — parada de línea sin pérdida de producto si se actúa rápido.",
        },
    ]

    for err in errores:
        doc.add_heading(f"{err['codigo']} — {err['titulo']}", level=1)

        p = doc.add_paragraph()
        p.add_run("Síntomas: ").bold = True
        p.add_run(err["sintomas"])

        doc.add_paragraph().add_run("Causas más frecuentes:").bold = True
        for causa in err["causas"]:
            doc.add_paragraph(causa, style="List Bullet")

        doc.add_paragraph().add_run("Acciones correctivas (por orden):").bold = True
        for i, accion in enumerate(err["acciones"], 1):
            doc.add_paragraph(f"{i}. {accion}")

        p2 = doc.add_paragraph()
        p2.add_run("Tiempo estimado de resolución: ").bold = True
        p2.add_run(err["tiempo_resolucion"])

        p3 = doc.add_paragraph()
        p3.add_run("Impacto en producción: ").bold = True
        p3.add_run(err["impacto"])
        doc.add_paragraph()

    path = DOCS_DIR / "Manual_Errores_Envasadora_UPS4.docx"
    doc.save(str(path))
    print(f"  ✓ Creado: {path.name}")


# ─────────────────────────────────────────────────────────────
# 3. Excel con códigos de error y soluciones
# ─────────────────────────────────────────────────────────────
def create_excel_codigos():
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Códigos de Error"

    header_fill   = PatternFill("solid", fgColor="003970")
    header_font   = Font(color="FFFFFF", bold=True, size=11)
    alt_fill      = PatternFill("solid", fgColor="EEF4FB")
    red_fill      = PatternFill("solid", fgColor="FDECEA")
    yellow_fill   = PatternFill("solid", fgColor="FFF8E1")
    green_fill    = PatternFill("solid", fgColor="E8F5E9")
    center_align  = Alignment(horizontal="center", vertical="center", wrap_text=True)
    wrap_align    = Alignment(vertical="top", wrap_text=True)

    headers = [
        "Código Error", "Descripción", "Equipo / Zona", "Severidad",
        "Causa Raíz Probable", "Acción Inmediata (Operario)",
        "Escalar a", "Tiempo Resolución Estimado", "Observaciones",
    ]
    ws.append(headers)
    for col, _ in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col)
        cell.fill   = header_fill
        cell.font   = header_font
        cell.alignment = center_align

    col_widths = [14, 32, 20, 12, 38, 42, 16, 22, 30]
    for i, w in enumerate(col_widths, 1):
        ws.column_dimensions[openpyxl.utils.get_column_letter(i)].width = w
    ws.row_dimensions[1].height = 36

    errores = [
        ("E-001", "Rotura lámina aluminio — zona sellado", "Zona sellado", "ALTO",
         "Temperatura sellado fuera de rango o lámina defectuosa",
         "1. Detener línea\n2. Verificar T° platos sellado (nominal 180°C±5°C)\n3. Inspeccionar rollo aluminio\n4. Reducir velocidad 10%",
         "Mantenimiento", "10–60 min", "Si persiste: ERR-001 en GMAO"),
        ("E-002", "Cavidades vacías en blíster — comprimido faltante", "Alimentación / Tolva", "ALTO",
         "Tolva vacía, vibrador apagado o guías obstruidas",
         "1. Verificar nivel tolva\n2. Activar vibrador desde HMI\n3. Limpiar guías con brocha",
         "Mantenimiento (si vibrador falla)", "5–15 min", "Sistema visión debe rechazar automáticamente"),
        ("E-003", "Atasco zona troquelado / corte blíster", "Troquelado", "MEDIO",
         "Desalineación de material o cuchilla desgastada",
         "1. Parada emergencia si riesgo\n2. Retirar material acumulado\n3. Realinear guía\n4. 3 ciclos en vacío antes de reiniciar",
         "Mantenimiento (cambio cuchilla)", "5–45 min", "Nunca extraer con manos en zona cuchillas"),
        ("E-004", "Temperatura formado PVC fuera de rango", "Zona formado", "ALTO",
         "Resistencia averiada, sensor descalibrado o PVC incorrecto",
         "1. Verificar T° en HMI (nominal 120°C±5°C)\n2. Si desviación >10°C más de 2 min: detener y avisar",
         "Mantenimiento urgente", "5 min diagnóstico / 1–3 h resolución", "No producir con T° fuera de rango"),
        ("E-005", "Impresión código Datamatrix/lote defectuosa", "Impresora inkjet", "CRÍTICO",
         "Tinta agotada, cabezal obstruido o parámetros incorrectos",
         "1. Verificar nivel tinta (<15% → cambiar)\n2. Ciclo limpieza cabezal\n3. Verificar parámetros lote/caducidad\n4. Prueba 5 unidades con lector",
         "Mantenimiento (fallo cabezal)", "10–60 min", "Sin código válido el producto NO puede salir"),
        ("E-006", "Exceso rechazo control de peso IPC", "IPC en línea", "CRÍTICO",
         "Variabilidad de comprimidos, báscula descalibrada o mezcla de lotes",
         "1. Detener llenado tolva\n2. Pesar muestra manual (10 uds)\n3. Verificar calibración báscula con patrones\n4. Segregar material afectado",
         "Control de Calidad", "15 min + investigación QC", "Blísteres afectados a cuarentena"),
        ("E-007", "Atasco / desbordamiento en encajadora", "Encajadora / Cartoning", "MEDIO",
         "Velocidad desincronizada, folleto mal insertado o cajas defectuosas",
         "1. Parar blistera y encajadora coordinadamente\n2. Retirar blísteres acumulados\n3. Verificar cargador folletos\n4. Reajustar velocidad sincronización HMI → Sync",
         "Mantenimiento (si sensor falla)", "10–20 min", "No forzar reinicio sin resolver causa"),
        ("E-008", "Alarma de temperatura sala — fuera de rango (+15°C a +25°C)", "Sala producción", "ALTO",
         "Fallo HVAC o exceso de calor generado por equipos",
         "1. Verificar consola HVAC en sala de control\n2. Si T° >27°C: notificar a Mantenimiento y Supervisor\n3. No iniciar producción si T° fuera de rango",
         "Mantenimiento + Supervisor", "30–90 min", "Registrar desviación ambiental en hoja de sala"),
        ("E-009", "Humedad relativa sala fuera de rango (RH 30–65%)", "Sala producción", "ALTO",
         "Fallo sistema humectación/deshumectación HVAC",
         "1. Verificar lectores de HR en sala (mínimo 2 puntos)\n2. Si RH>65% con comprimidos higroscópicos: detener producción\n3. Notificar a QC y Mantenimiento",
         "QC + Mantenimiento", "30–120 min", "Impacto en estabilidad del producto"),
        ("E-010", "Fallo sensor de presencia — zona de transferencia blíster", "Sensores / PLC", "MEDIO",
         "Suciedad en sensor óptico o cable dañado",
         "1. Limpiar lente del sensor con paño seco\n2. Verificar alineación del haz\n3. Si persiste, registrar en GMAO y avisar a Mantenimiento",
         "Mantenimiento (electricidad)", "5–30 min", "No puentear sensores en producción"),
        ("E-011", "Ruido anormal en zona de transmisión / cadena", "Mecánica transmisión", "MEDIO",
         "Cadena desengrasada, piñón desgastado o cuerpo extraño",
         "1. Detener línea inmediatamente (daño potencial grave)\n2. Inspeccionar visualmente zona de ruido\n3. Avisar a Mantenimiento",
         "Mantenimiento urgente", "30–120 min", "No reiniciar hasta resolución por Mantenimiento"),
        ("E-012", "Alarma PLC — fallo comunicación entre módulos", "Sistema de control (PLC)", "ALTO",
         "Cable Ethernet/Profibus dañado, módulo E/S averiado o software error",
         "1. Anotar el código exacto de alarma en pantalla HMI\n2. Intentar reinicio suave desde HMI (Reset Soft)\n3. Si persiste, avisar a Automatización",
         "Automatización / IT", "30 min – 4 horas", "No apagar PLC principal sin autorización"),
        ("E-013", "Película PVC con burbujas o deformaciones visibles", "Zona formado", "MEDIO",
         "Humedad en el rollo de PVC o temperatura de formado demasiado alta",
         "1. Retirar el rollo de PVC y verificar almacenamiento (debe estar <25°C, <65%HR)\n2. Sustituir por rollo nuevo del almacén\n3. Realizar ajuste fino de temperatura (-5°C) si persiste",
         "Almacén (rollo nuevo) / QC (si hay duda de material)", "15–30 min", "Conservar rollo afectado para investigación"),
        ("E-014", "Sistema de serialización / Track & Trace — fallo de conexión", "Serialización", "CRÍTICO",
         "Fallo de red, servidor S&T no disponible o licencia caducada",
         "1. Verificar conectividad red en PC de serialización (ping al servidor)\n2. Notificar a IT inmediatamente\n3. No producir sin serialización activa",
         "IT / Serialización", "Variable — puede requerir escalada a proveedor", "Cumplimiento regulatorio: es obligatorio"),
        ("E-015", "Blísteres con lámina arrugada tras sellado", "Zona sellado", "BAJO",
         "Tensión de bobina de aluminio incorrecta o temperatura desigual en platos",
         "1. Ajustar freno de bobina de aluminio (aumentar ligeramente)\n2. Verificar distribución uniforme de temperatura en plato de sellado con termómetro de contacto\n3. Si la arruga es sistemática, avisar a Mantenimiento",
         "Mantenimiento (si ajuste no resuelve)", "10–20 min", "Impacto estético — verificar si supera límite de aceptación visual"),
    ]

    severity_color = {"CRÍTICO": red_fill, "ALTO": red_fill,
                      "MEDIO": yellow_fill, "BAJO": green_fill}

    for i, row_data in enumerate(errores, 2):
        ws.append(list(row_data))
        for col in range(1, len(headers) + 1):
            cell = ws.cell(row=i, column=col)
            cell.alignment = wrap_align if col > 1 else center_align
            if col == 4:
                cell.fill = severity_color.get(row_data[3], PatternFill())
                cell.font = Font(bold=True)
                cell.alignment = center_align
            elif i % 2 == 0:
                cell.fill = alt_fill
        ws.row_dimensions[i].height = 75

    # Segunda hoja: Checklist de inicio de turno
    ws2 = wb.create_sheet("Checklist Inicio Turno")
    ws2.column_dimensions["A"].width = 5
    ws2.column_dimensions["B"].width = 55
    ws2.column_dimensions["C"].width = 15
    ws2.column_dimensions["D"].width = 30

    ws2["A1"] = "CHECKLIST DE VERIFICACIÓN — INICIO DE TURNO"
    ws2["A1"].font = Font(bold=True, size=13, color="003970")
    ws2.merge_cells("A1:D1")
    ws2["A1"].alignment = center_align

    checklist = [
        ("SEGURIDAD", [
            "EPIs disponibles y en buen estado (guantes, mascarilla, gafas)",
            "Extintor en posición y sin señal de uso",
            "Salidas de emergencia despejadas",
            "Ficha de seguridad del producto disponible en línea",
        ]),
        ("LIMPIEZA", [
            "Registro de limpieza del turno anterior firmado",
            "Ausencia visible de restos de producto anterior",
            "Papeleras y contenedores de rechazo vacios",
        ]),
        ("EQUIPOS", [
            "Tolva de alimentación llena con producto del lote correcto",
            "Rollo de PVC cargado correctamente",
            "Rollo de aluminio cargado correctamente",
            "Cargador de folletos lleno",
            "Cargador de cajas lleno",
            "Impresora inkjet con tinta suficiente (>30%)",
            "Temperatura sala OK (15–25°C) — verificar display",
            "Humedad sala OK (30–65%HR) — verificar display",
        ]),
        ("DOCUMENTACIÓN", [
            "Orden de fabricación / packaging order disponible y firmada",
            "Parámetros de formato configurados en HMI según la orden",
            "Sistema de serialización activo y comunicado con servidor",
            "Registro de inicio de turno cumplimentado",
        ]),
    ]

    current_row = 3
    for seccion, items in checklist:
        ws2.cell(row=current_row, column=1, value=seccion).font = Font(bold=True, color="FFFFFF")
        ws2.cell(row=current_row, column=1).fill = PatternFill("solid", fgColor="003970")
        ws2.merge_cells(f"A{current_row}:D{current_row}")
        ws2.cell(row=current_row, column=1).alignment = center_align
        current_row += 1
        for item in items:
            ws2.cell(row=current_row, column=1, value="☐").alignment = center_align
            ws2.cell(row=current_row, column=2, value=item).alignment = wrap_align
            ws2.cell(row=current_row, column=3, value="OK / NOK").font = Font(italic=True, color="888888")
            ws2.cell(row=current_row, column=3).alignment = center_align
            ws2.cell(row=current_row, column=4).alignment = wrap_align
            ws2.row_dimensions[current_row].height = 20
            current_row += 1
        current_row += 1

    path = DOCS_DIR / "Codigos_Error_Soluciones.xlsx"
    wb.save(str(path))
    print(f"  ✓ Creado: {path.name}")


# ─────────────────────────────────────────────────────────────
# Main
# ─────────────────────────────────────────────────────────────
if __name__ == "__main__":
    print("Generando documentos de ejemplo en docs/...")
    try:
        create_sop_limpieza()
        create_manual_errores()
        create_excel_codigos()
        print("\n✓ Documentos creados correctamente.")
        print("  Ejecuta ahora:  python ingest.py  para indexarlos.")
    except ImportError as e:
        print(f"\n✗ Dependencia faltante: {e}")
        print("  Instala con:  pip install python-docx openpyxl")
