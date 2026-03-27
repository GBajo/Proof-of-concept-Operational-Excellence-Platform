"""
ingest.py — Módulo de ingesta de documentos para OpEx Platform RAG

Uso desde línea de comandos:
    python ingest.py                    # indexa todo docs/
    python ingest.py --file docs/sop.pdf
    python ingest.py --url https://...
    python ingest.py --clear            # borra la base y re-indexa

Uso desde Python / Flask:
    from ingest import run_ingest
    result = run_ingest()               # devuelve dict con estadísticas
"""

from __future__ import annotations

import argparse
import os
import re
import sqlite3
import sys
from pathlib import Path
from typing import Iterator

# ── Constantes ──────────────────────────────────────────────
DOCS_DIR    = Path(__file__).parent / "docs"
DB_PATH     = os.environ.get("DATABASE_PATH", "opex.db")
CHUNK_WORDS = 500
OVERLAP_WORDS = 50

# ── Chunking ─────────────────────────────────────────────────

def chunk_text(text: str, chunk_words: int = CHUNK_WORDS,
               overlap: int = OVERLAP_WORDS) -> list[str]:
    """Divide texto en fragmentos de ~chunk_words palabras con solapamiento."""
    words = text.split()
    chunks: list[str] = []
    start = 0
    while start < len(words):
        end = start + chunk_words
        chunks.append(" ".join(words[start:end]))
        start += chunk_words - overlap
    return [c for c in chunks if len(c.strip()) > 20]


# ── Extractores ──────────────────────────────────────────────

def extract_pdf(path: Path) -> str:
    """Extrae texto de un PDF con pdfplumber (fallback: PyPDF2)."""
    try:
        import pdfplumber
        with pdfplumber.open(path) as pdf:
            return "\n".join(
                page.extract_text() or "" for page in pdf.pages
            )
    except ImportError:
        pass
    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(str(path))
        return "\n".join(
            page.extract_text() or "" for page in reader.pages
        )
    except ImportError:
        raise RuntimeError(
            "Instala pdfplumber o PyPDF2:  pip install pdfplumber"
        )


def extract_docx(path: Path) -> str:
    """Extrae texto de un archivo Word (.docx)."""
    try:
        from docx import Document
    except ImportError:
        raise RuntimeError("Instala python-docx:  pip install python-docx")
    doc = Document(str(path))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    # Incluir también tablas
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))
    return "\n".join(paragraphs)


def extract_xlsx(path: Path) -> str:
    """Extrae texto de un Excel. Cada fila se convierte en un fragmento."""
    try:
        import openpyxl
    except ImportError:
        raise RuntimeError("Instala openpyxl:  pip install openpyxl")
    wb = openpyxl.load_workbook(str(path), read_only=True, data_only=True)
    rows_text: list[str] = []
    for sheet in wb.worksheets:
        headers: list[str] = []
        for i, row in enumerate(sheet.iter_rows(values_only=True)):
            cells = [str(c) if c is not None else "" for c in row]
            if i == 0:
                headers = cells
                continue
            if all(c == "" for c in cells):
                continue
            if headers:
                row_str = " | ".join(
                    f"{h}: {v}" for h, v in zip(headers, cells) if v
                )
            else:
                row_str = " | ".join(c for c in cells if c)
            if row_str.strip():
                rows_text.append(row_str)
    wb.close()
    return "\n".join(rows_text)


def extract_url(url: str) -> str:
    """Descarga una URL y extrae el texto visible."""
    try:
        import requests
        from bs4 import BeautifulSoup
    except ImportError:
        raise RuntimeError(
            "Instala requests y beautifulsoup4:  "
            "pip install requests beautifulsoup4"
        )
    resp = requests.get(url, timeout=15)
    resp.raise_for_status()
    soup = BeautifulSoup(resp.text, "html.parser")
    # Eliminar scripts y estilos
    for tag in soup(["script", "style", "nav", "footer", "header"]):
        tag.decompose()
    return soup.get_text(separator="\n")


# ── Detección de tipo ─────────────────────────────────────────

def detect_type(path: Path) -> str:
    ext = path.suffix.lower()
    return {".pdf": "pdf", ".docx": "docx",
            ".xlsx": "xlsx", ".xls": "xlsx"}.get(ext, "")


def extract(path: Path) -> tuple[str, str]:
    """Devuelve (texto, source_type)."""
    t = detect_type(path)
    if t == "pdf":
        return extract_pdf(path), "pdf"
    if t == "docx":
        return extract_docx(path), "docx"
    if t == "xlsx":
        return extract_xlsx(path), "xlsx"
    raise ValueError(f"Tipo de archivo no soportado: {path.suffix}")


# ── Base de datos ─────────────────────────────────────────────

def get_conn(db_path: str = DB_PATH) -> sqlite3.Connection:
    conn = sqlite3.connect(db_path)
    conn.execute("PRAGMA foreign_keys = ON")
    conn.execute("""
        CREATE TABLE IF NOT EXISTS knowledge_base (
            id          INTEGER PRIMARY KEY AUTOINCREMENT,
            source_file TEXT NOT NULL,
            source_type TEXT NOT NULL DEFAULT 'file',
            chunk_index INTEGER NOT NULL,
            chunk_text  TEXT NOT NULL,
            indexed_at  TEXT NOT NULL DEFAULT (datetime('now'))
        )
    """)
    conn.execute(
        "CREATE INDEX IF NOT EXISTS idx_kb_source "
        "ON knowledge_base(source_file)"
    )
    conn.commit()
    return conn


def delete_source(conn: sqlite3.Connection, source_file: str) -> None:
    conn.execute(
        "DELETE FROM knowledge_base WHERE source_file = ?", (source_file,)
    )
    conn.commit()


def insert_chunks(conn: sqlite3.Connection, source_file: str,
                  source_type: str, chunks: list[str]) -> None:
    conn.executemany(
        """INSERT INTO knowledge_base
           (source_file, source_type, chunk_index, chunk_text)
           VALUES (?, ?, ?, ?)""",
        [(source_file, source_type, i, chunk)
         for i, chunk in enumerate(chunks)],
    )
    conn.commit()


# ── Búsqueda (para el módulo RAG) ────────────────────────────

def search(query: str, top_k: int = 5,
           db_path: str = DB_PATH) -> list[dict]:
    """
    Búsqueda por palabras clave en knowledge_base.
    Devuelve los top_k fragmentos más relevantes.
    """
    conn = sqlite3.connect(db_path)
    conn.row_factory = sqlite3.Row
    words = [w.strip() for w in query.lower().split() if len(w) > 2]
    if not words:
        return []

    # Puntuación: número de palabras del query que aparecen en el chunk
    rows = conn.execute(
        "SELECT id, source_file, source_type, chunk_index, chunk_text "
        "FROM knowledge_base"
    ).fetchall()
    conn.close()

    scored: list[tuple[int, dict]] = []
    for row in rows:
        text_lower = row["chunk_text"].lower()
        score = sum(1 for w in words if w in text_lower)
        if score > 0:
            scored.append((score, dict(row)))

    scored.sort(key=lambda x: x[0], reverse=True)
    return [item for _, item in scored[:top_k]]


# ── Lógica principal ──────────────────────────────────────────

def index_file(path: Path, conn: sqlite3.Connection,
               force: bool = False) -> dict:
    """Indexa un único archivo. Devuelve estadísticas."""
    source_name = path.name
    # Comprobar si ya está indexado
    existing = conn.execute(
        "SELECT COUNT(*) FROM knowledge_base WHERE source_file = ?",
        (source_name,),
    ).fetchone()[0]

    if existing and not force:
        return {"file": source_name, "status": "skipped",
                "chunks": existing}

    try:
        text, src_type = extract(path)
        text = re.sub(r"\s{3,}", "\n\n", text).strip()
        if not text:
            return {"file": source_name, "status": "empty", "chunks": 0}
        chunks = chunk_text(text)
        delete_source(conn, source_name)
        insert_chunks(conn, source_name, src_type, chunks)
        return {"file": source_name, "status": "ok",
                "chunks": len(chunks), "type": src_type}
    except Exception as e:
        return {"file": source_name, "status": "error", "error": str(e)}


def index_url(url: str, conn: sqlite3.Connection,
              force: bool = False) -> dict:
    """Indexa una URL."""
    source_name = url
    existing = conn.execute(
        "SELECT COUNT(*) FROM knowledge_base WHERE source_file = ?",
        (source_name,),
    ).fetchone()[0]

    if existing and not force:
        return {"file": source_name, "status": "skipped",
                "chunks": existing}
    try:
        text = extract_url(url)
        text = re.sub(r"\s{3,}", "\n\n", text).strip()
        chunks = chunk_text(text)
        delete_source(conn, source_name)
        insert_chunks(conn, source_name, "url", chunks)
        return {"file": source_name, "status": "ok",
                "chunks": len(chunks), "type": "url"}
    except Exception as e:
        return {"file": source_name, "status": "error", "error": str(e)}


def run_ingest(docs_dir: Path = DOCS_DIR, db_path: str = DB_PATH,
               force: bool = False) -> dict:
    """
    Indexa todos los documentos de docs_dir.
    Devuelve un resumen con estadísticas.
    """
    docs_dir.mkdir(exist_ok=True)
    conn = get_conn(db_path)
    results: list[dict] = []

    supported = {".pdf", ".docx", ".xlsx", ".xls"}
    files = [f for f in sorted(docs_dir.iterdir())
             if f.is_file() and f.suffix.lower() in supported]

    for path in files:
        result = index_file(path, conn, force=force)
        results.append(result)
        status = result["status"]
        icon = "✓" if status == "ok" else "↷" if status == "skipped" else "✗"
        print(f"  {icon} {result['file']} — {status}"
              + (f" ({result.get('chunks', 0)} fragmentos)" if "chunks" in result else "")
              + (f": {result.get('error','')}" if status == "error" else ""))

    conn.close()
    ok      = sum(1 for r in results if r["status"] == "ok")
    skipped = sum(1 for r in results if r["status"] == "skipped")
    errors  = sum(1 for r in results if r["status"] == "error")
    total_chunks = sum(r.get("chunks", 0)
                       for r in results if r["status"] == "ok")

    return {
        "total_files": len(results),
        "ok": ok,
        "skipped": skipped,
        "errors": errors,
        "total_chunks": total_chunks,
        "details": results,
    }


# ── CLI ───────────────────────────────────────────────────────

def main() -> None:
    parser = argparse.ArgumentParser(
        description="OpEx Platform — Indexador de documentos para RAG"
    )
    parser.add_argument("--file",  help="Indexar un archivo concreto")
    parser.add_argument("--url",   help="Indexar una URL")
    parser.add_argument("--clear", action="store_true",
                        help="Borrar toda la base de conocimiento antes de indexar")
    parser.add_argument("--force", action="store_true",
                        help="Re-indexar aunque el archivo ya esté indexado")
    parser.add_argument("--db",    default=DB_PATH,
                        help="Ruta a la base de datos SQLite")
    args = parser.parse_args()

    conn = get_conn(args.db)

    if args.clear:
        conn.execute("DELETE FROM knowledge_base")
        conn.commit()
        print("Base de conocimiento vaciada.")

    if args.file:
        path = Path(args.file)
        if not path.exists():
            print(f"ERROR: archivo no encontrado: {path}", file=sys.stderr)
            sys.exit(1)
        r = index_file(path, conn, force=True)
        print(f"  {r['status']}: {r['file']} — {r.get('chunks', 0)} fragmentos")
    elif args.url:
        r = index_url(args.url, conn, force=True)
        print(f"  {r['status']}: {r['file']} — {r.get('chunks', 0)} fragmentos")
    else:
        print(f"Indexando documentos en: {DOCS_DIR}")
        summary = run_ingest(force=args.force, db_path=args.db)
        print(f"\nResumen: {summary['ok']} indexados, "
              f"{summary['skipped']} omitidos, "
              f"{summary['errors']} errores, "
              f"{summary['total_chunks']} fragmentos totales.")

    conn.close()


if __name__ == "__main__":
    main()
